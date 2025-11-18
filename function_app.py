import azure.functions as func
import logging
import io
import os
import base64
from typing import Dict, List, Optional
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from openai import AzureOpenAI
import json
from PIL import Image
import re

app = func.FunctionApp()


def apply_text_formatting(paragraph, text: str):
    """
    Aplica formatações especiais ao texto do parágrafo.
    - *palavra* -> itálico
    - **palavra** -> negrito (se vier do GPT)
    - <<ALT_CORRETA_INICIO>> ... <<ALT_CORRETA_FIM>> -> negrito
    
    Args:
        paragraph: Objeto Paragraph do python-docx
        text: Texto com marcadores de formatação
    """
    # Limpar runs existentes
    for run in paragraph.runs:
        run.text = ""
    
    # Processar marcadores de alternativa correta
    # <<ALT_CORRETA_INICIO>> texto <<ALT_CORRETA_FIM>> -> negrito
    alt_pattern = r'<<ALT_CORRETA_INICIO>>(.+?)<<ALT_CORRETA_FIM>>'
    if '<<ALT_CORRETA_INICIO>>' in text:
        parts = re.split(alt_pattern, text, flags=re.DOTALL)
        for i, part in enumerate(parts):
            if i % 2 == 1:  # Parte entre os marcadores
                run = paragraph.add_run(part)
                run.bold = True
            else:
                # Processar itálicos na parte normal
                apply_italic_formatting(paragraph, part)
    else:
        # Processar itálicos
        apply_italic_formatting(paragraph, text)


def apply_italic_formatting(paragraph, text: str):
    """
    Aplica formatação de itálico (*palavra*).
    
    Args:
        paragraph: Objeto Paragraph do python-docx
        text: Texto com marcadores de itálico
    """
    # Padrão para itálico: *palavra* (mas não **palavra**)
    italic_pattern = r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)'
    parts = re.split(italic_pattern, text)
    
    for i, part in enumerate(parts):
        if i % 2 == 1:  # Parte entre asteriscos simples
            run = paragraph.add_run(part)
            run.italic = True
        elif part:  # Parte normal
            paragraph.add_run(part)

# Configuração do Azure OpenAI
AZURE_OPENAI_ENDPOINT = os.environ.get("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_API_KEY = os.environ.get("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_DEPLOYMENT = os.environ.get("AZURE_OPENAI_DEPLOYMENT", "gpt-4")
AZURE_OPENAI_API_VERSION = os.environ.get("AZURE_OPENAI_API_VERSION", "2024-02-15-preview")

# Inicializar cliente OpenAI
client = AzureOpenAI(
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION
)


def describe_image(image_bytes: bytes, context: str = "") -> str:
    """
    Gera descrição pedagógica de imagem usando Azure OpenAI Vision (GPT-4o).
    Descrição será inserida no texto do documento após a imagem.
    
    Args:
        image_bytes: Bytes da imagem
        context: Contexto adicional sobre a imagem (opcional)
        
    Returns:
        Descrição pedagógica da imagem em português
    """
    try:
        # Converter imagem para base64
        base64_image = base64.b64encode(image_bytes).decode('utf-8')
        
        # Determinar tipo MIME da imagem
        try:
            img = Image.open(io.BytesIO(image_bytes))
            img_format = img.format.lower()
            mime_type = f"image/{img_format}" if img_format in ['jpeg', 'jpg', 'png', 'gif', 'webp'] else "image/jpeg"
        except:
            mime_type = "image/jpeg"
        
        system_prompt = """Você é um revisor pedagógico do SENAC/SC especializado em descrição de imagens.

OBJETIVO:
Descrever a imagem de forma DIDÁTICA, CLARA e DETALHADA, como se estivesse explicando para um aluno.
Use linguagem dialógica e explicativa, transformando elementos visuais em texto compreensível.

REGRAS:
1. Use linguagem clara e acessível, sem jargões técnicos não explicados
2. Descreva TODOS os elementos relevantes: gráficos, tabelas, diagramas, textos visíveis, cores, formas
3. Se for gráfico/tabela: descreva os dados, tendências, valores principais
4. Se for diagrama/fluxo: explique o processo, conexões, etapas
5. Se for foto/ilustração: descreva cenário, pessoas, objetos, ações
6. Se houver texto na imagem: transcreva-o integralmente
7. Organize a descrição de forma lógica (do geral ao específico)
8. Use tom explicativo e pedagógico

FORMATO DA RESPOSTA:
Inicie sempre com "Descrição da imagem:" seguido da descrição completa em português.
Seja detalhado mas objetivo. Mínimo 2 parágrafos, máximo 5 parágrafos."""

        user_prompt = "Descreva detalhadamente esta imagem de forma pedagógica e didática."
        if context:
            user_prompt += f"\n\nContexto do documento: {context}"
        
        response = client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            messages=[
                {"role": "system", "content": system_prompt},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": user_prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{base64_image}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=1500,
            temperature=0.3
        )
        
        description = response.choices[0].message.content.strip()
        logging.info(f"✅ Imagem descrita: {description[:80]}...")
        return description
        
    except Exception as e:
        logging.error(f"Erro ao descrever imagem: {str(e)}")
        return "Descrição da imagem: Imagem sem descrição disponível devido a erro técnico."


def process_paragraph_text(text: str, is_table_cell: bool = False) -> str:
    """
    Processa um parágrafo usando Azure OpenAI com revisão pedagógica SENAC.
    
    Args:
        text: Texto do parágrafo a ser revisado
        is_table_cell: Se True, aplica processamento específico para células de tabela
        
    Returns:
        Texto revisado pedagogicamente
    """
    if not text or len(text.strip()) == 0:
        return text
    
    # Detectar e preservar tokens de mídia
    import re
    media_tokens = re.findall(r'\[\[(FIG|TAB|SA)\d+\]\]', text)
    
    try:
        system_prompt = """Você é revisor pedagógico do SENAC/SC.

OBJETIVO:
Entregar o texto revisado, didático e padronizado, pronto para publicação.
O texto deve soar como uma AULA, em tom explicativo e próximo ao aluno, quase como uma conversa.
Use linguagem dialógica e interações leves ("Você sabia…?", "Reflita…", "Agora pense…", "Vamos entender…") para engajar o aluno.
Devolva exclusivamente o texto revisado, sem qualquer comentário, explicação, preâmbulo ou cabeçalho extra.

REGRAS OBRIGATÓRIAS:
0) PROIBIDO qualquer meta-texto/comentário fora do conteúdo (ex.: "Segue o texto...", "O texto foi revisado...").
0a) PROIBIDO inserir placeholders como "..." ou "(continua...)". NUNCA encerre com frase incompleta.
1) MODO CÓPIA MELHORADA: mantenha as frases próximas do original. Corrija ortografia, gramática, pontuação, concordância e coesão.
   PORÉM, MELHORE a linguagem para ser mais dialógica e pedagógica, sem reescrita total.
2) Simplifique linguagem técnica mantendo precisão. Explique termos complexos em linguagem acessível.
3) Use TOM CONVERSACIONAL como em aula: 1ª pessoa do plural ("vamos", "veremos"), perguntas retóricas, interações.
4) PARÁGRAFOS CURTOS: divida parágrafos longos em parágrafos menores (máximo 5-6 linhas cada).
5) FRASES CLARAS: divida frases muito longas em frases mais curtas e diretas.
6) INSIRA nomes fictícios para empresas, pessoas, instituições quando aplicável (ex: "Empresa TechSolutions", "João Silva").
   Mantenha o MESMO nome fictício em todo o texto.
7) Preserve estrutura, ordem, exemplos, tabelas, listas.
8) Padronize títulos/subtítulos em CAIXA ALTA quando forem cabeçalhos principais.
9) TERMOS TÉCNICOS: simplifique ou explique brevemente quando aparecerem pela primeira vez.
10) PALAVRAS ESTRANGEIRAS: coloque em itálico (retorne com marcador *palavra* para indicar itálico).
11) REMOVA linguagem excessivamente formal ou acadêmica.
12) ADICIONE pequenos elementos pedagógicos quando natural: "Observe que...", "Note que...", "É importante destacar...".
13) TOKENS DE MÍDIA ([[FIG1]], [[TAB1]], [[SA1]]): PRESERVE EXATAMENTE onde estão. NUNCA remova, renomeie ou mova.
14) NÃO remova citações, autores, anos, referências bibliográficas.
15) MANTENHA o comprimento similar ao original - não resuma nem encurte drasticamente.
16) NÃO use markdown (##, **, __, ---).
17) ALTERNATIVAS DE QUESTÕES: se detectar questões de múltipla escolha, identifique a alternativa correta e envolva
    APENAS A LINHA DA ALTERNATIVA com <<ALT_CORRETA_INICIO>> texto da alternativa <<ALT_CORRETA_FIM>>.

IMPORTANTE: Retorne SOMENTE o texto revisado. Sem comentários, sem explicações, sem preâmbulos."""

        response = client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"TEXTO ORIGINAL:\n{text}"}
            ],
            temperature=0.4,  # Aumentado para permitir mais criatividade pedagógica
            max_tokens=6000
        )
        
        corrected_text = response.choices[0].message.content.strip()
        
        # Garantir que tokens de mídia foram preservados
        for token in media_tokens:
            if token not in corrected_text:
                logging.warning(f"Token {token} foi removido, restaurando...")
                # Tentar restaurar o token
                corrected_text = text  # Fallback para texto original se tokens forem removidos
                break
        
        return corrected_text
        
    except Exception as e:
        logging.error(f"Erro ao processar parágrafo com OpenAI: {str(e)}")
        return text  # Retorna texto original em caso de erro


def process_word_document(file_content: bytes, describe_images: bool = True) -> bytes:
    """
    Processa documento Word completo mantendo formatação, imagens, tabelas, etc.
    Adiciona descrições automáticas às imagens usando Azure OpenAI Vision.
    
    Args:
        file_content: Conteúdo binário do documento Word
        describe_images: Se True, adiciona descrições às imagens
        
    Returns:
        Conteúdo binário do documento corrigido
    """
    # Carregar documento da memória
    doc_stream = io.BytesIO(file_content)
    doc = Document(doc_stream)
    
    logging.info(f"Processando documento com {len(doc.paragraphs)} parágrafos")
    
    # Contar imagens no documento
    images_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            images_count += 1
    logging.info(f"Imagens encontradas no documento: {images_count}")
    
    # Processar cada parágrafo mantendo formatação
    paragraphs_processed = 0
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Apenas processar parágrafos com texto
            try:
                original_text = paragraph.text
                corrected_text = process_paragraph_text(original_text, is_table_cell=False)
                
                # Aplicar formatações especiais e preservar estilo
                if corrected_text != original_text:
                    # Aplicar formatações (itálico, negrito, marcadores)
                    apply_text_formatting(paragraph, corrected_text)
                    paragraphs_processed += 1
                    
            except Exception as e:
                logging.error(f"Erro ao processar parágrafo: {str(e)}")
                continue
    
    logging.info(f"Total de parágrafos corrigidos: {paragraphs_processed}")
    
    # Processar e descrever imagens
    if describe_images and images_count > 0:
        logging.info("🖼️ Iniciando descrição de imagens...")
        images_described = 0
        
        try:
            # Precisamos iterar em ordem reversa para não afetar índices ao inserir parágrafos
            paragraphs_list = list(doc.paragraphs)
            
            for para_idx in range(len(paragraphs_list) - 1, -1, -1):
                paragraph = paragraphs_list[para_idx]
                
                for run in paragraph.runs:
                    # Verificar se o run contém imagem
                    if 'graphic' in run._element.xml:
                        try:
                            # Extrair a imagem
                            blip_elements = run._element.xpath('.//a:blip')
                            if blip_elements:
                                for blip in blip_elements:
                                    embed = blip.get(qn('r:embed'))
                                    if embed:
                                        image_part = doc.part.related_parts[embed]
                                        image_bytes = image_part.blob
                                        
                                        # Gerar descrição pedagógica
                                        # Buscar contexto dos parágrafos vizinhos
                                        context_parts = []
                                        if para_idx > 0:
                                            context_parts.append(paragraphs_list[para_idx - 1].text[:150])
                                        context_parts.append(paragraph.text[:150])
                                        if para_idx < len(paragraphs_list) - 1:
                                            context_parts.append(paragraphs_list[para_idx + 1].text[:150])
                                        context = " ".join(context_parts)
                                        
                                        description = describe_image(image_bytes, context)
                                        
                                        # Inserir descrição como NOVO PARÁGRAFO após a imagem
                                        # Encontrar o elemento do parágrafo no XML
                                        para_element = paragraph._element
                                        parent_element = para_element.getparent()
                                        
                                        # Criar novo parágrafo com a descrição
                                        new_para = doc.add_paragraph()
                                        new_para.text = description
                                        new_para_element = new_para._element
                                        
                                        # Inserir o novo parágrafo logo após o parágrafo da imagem
                                        parent_element.insert(
                                            parent_element.index(para_element) + 1,
                                            new_para_element
                                        )
                                        
                                        images_described += 1
                                        logging.info(f"  ✅ Imagem {images_described} descrita e inserida no texto")
                        except Exception as e:
                            logging.warning(f"  ⚠️ Erro ao processar imagem inline: {str(e)}")
            
            logging.info(f"✅ Total de imagens descritas: {images_described}")
            
        except Exception as e:
            logging.error(f"Erro ao processar imagens: {str(e)}")
    
    # Processar tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        try:
                            original_text = paragraph.text
                            corrected_text = process_paragraph_text(original_text, is_table_cell=True)
                            
                            if corrected_text != original_text:
                                apply_text_formatting(paragraph, corrected_text)
                                    
                        except Exception as e:
                            logging.error(f"Erro ao processar célula de tabela: {str(e)}")
                            continue
    
    # Salvar documento processado em memória
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    
    return output_stream.getvalue()


@app.route(route="correct-document", auth_level=func.AuthLevel.FUNCTION, methods=["POST"])
def correct_document(req: func.HttpRequest) -> func.HttpResponse:
    """
    Azure Function HTTP endpoint para correção de documentos Word.
    
    Endpoint: POST /api/correct-document
    Content-Type: multipart/form-data
    
    Parâmetros:
        - file: Arquivo .docx para correção (upload)
        
    Retorna:
        - Arquivo .docx corrigido
    """
    logging.info('Recebida requisição para correção de documento Word')
    
    try:
        # Validar variáveis de ambiente
        if not AZURE_OPENAI_ENDPOINT or not AZURE_OPENAI_API_KEY:
            return func.HttpResponse(
                json.dumps({
                    "error": "Configuração do Azure OpenAI não encontrada. Verifique as variáveis de ambiente."
                }),
                status_code=500,
                mimetype="application/json"
            )
        
        # Obter arquivo do request
        file = req.files.get('file')
        
        if not file:
            return func.HttpResponse(
                json.dumps({
                    "error": "Nenhum arquivo foi enviado. Use o campo 'file' no multipart/form-data"
                }),
                status_code=400,
                mimetype="application/json"
            )
        
        # Verificar extensão do arquivo
        filename = file.filename
        if not filename.lower().endswith('.docx'):
            return func.HttpResponse(
                json.dumps({
                    "error": "Apenas arquivos .docx são suportados"
                }),
                status_code=400,
                mimetype="application/json"
            )
        
        # Ler conteúdo do arquivo
        file_content = file.read()
        logging.info(f"Arquivo recebido: {filename} ({len(file_content)} bytes)")
        
        # Processar documento
        corrected_content = process_word_document(file_content)
        logging.info(f"Documento processado com sucesso ({len(corrected_content)} bytes)")
        
        # Retornar arquivo corrigido
        corrected_filename = filename.replace('.docx', '_corrigido.docx')
        
        return func.HttpResponse(
            body=corrected_content,
            status_code=200,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{corrected_filename}"'
            }
        )
        
    except Exception as e:
        logging.error(f"Erro ao processar documento: {str(e)}", exc_info=True)
        return func.HttpResponse(
            json.dumps({
                "error": f"Erro ao processar documento: {str(e)}"
            }),
            status_code=500,
            mimetype="application/json"
        )


@app.route(route="health", auth_level=func.AuthLevel.ANONYMOUS, methods=["GET"])
def health_check(req: func.HttpRequest) -> func.HttpResponse:
    """
    Endpoint de health check para verificar status da função.
    """
    return func.HttpResponse(
        json.dumps({
            "status": "healthy",
            "service": "word-correction-function",
            "azure_openai_configured": bool(AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_API_KEY)
        }),
        status_code=200,
        mimetype="application/json"
    )


@app.blob_trigger(arg_name="inputblob", 
                  path="documentos/input/{name}",
                  connection="AzureWebJobsStorage")
@app.blob_output(arg_name="outputblob",
                 path="documentos/output/{name}",
                 connection="AzureWebJobsStorage")
def blob_correct_document(inputblob: func.InputStream, outputblob: func.Out[bytes]) -> None:
    """
    Blob Trigger: Processa automaticamente documentos Word quando carregados no container.
    
    Trigger: Blob Upload em 'documentos/input/'
    Output: Blob em 'documentos/output/' com documento corrigido
    
    Exemplo:
    - Upload: documentos/input/documento.docx
    - Output: documentos/output/documento.docx (corrigido)
    """
    logging.info(f'🔔 Blob Trigger ativado!')
    logging.info(f'📄 Processando blob: {inputblob.name}')
    logging.info(f'📊 Tamanho: {inputblob.length} bytes')
    
    # Verificar se é um arquivo .docx
    if not inputblob.name.lower().endswith('.docx'):
        logging.warning(f'⚠️ Arquivo ignorado (não é .docx): {inputblob.name}')
        return
    
    try:
        # Validar configuração do Azure OpenAI
        if not AZURE_OPENAI_ENDPOINT or not AZURE_OPENAI_API_KEY:
            logging.error("❌ Azure OpenAI não configurado!")
            return
        
        # Ler conteúdo do blob
        file_content = inputblob.read()
        logging.info(f'✅ Arquivo lido: {len(file_content)} bytes')
        
        # Processar documento
        logging.info('⚙️ Iniciando processamento com Azure OpenAI...')
        corrected_content = process_word_document(file_content)
        
        # Escrever no blob de saída
        outputblob.set(corrected_content)
        
        logging.info(f'✅ Documento processado com sucesso!')
        logging.info(f'📤 Salvo em: documentos/output/{inputblob.name.split("/")[-1]}')
        logging.info(f'📊 Tamanho final: {len(corrected_content)} bytes')
        
    except Exception as e:
        logging.error(f'❌ Erro ao processar {inputblob.name}: {str(e)}', exc_info=True)