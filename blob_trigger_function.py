"""
Azure Function com Blob Storage Trigger para processamento em lote.

Esta versão alternativa monitora um container de Blob Storage e processa
automaticamente documentos Word que são carregados.

Arquitetura:
1. Upload do documento → Blob Storage (container: input-documents)
2. Blob Trigger → Azure Function (processamento)
3. Documento corrigido → Blob Storage (container: corrected-documents)
"""

import azure.functions as func
import logging
import io
import os
from docx import Document
from openai import AzureOpenAI
from datetime import datetime

# Configuração do Azure OpenAI
AZURE_OPENAI_ENDPOINT = os.environ.get("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_API_KEY = os.environ.get("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_DEPLOYMENT = os.environ.get("AZURE_OPENAI_DEPLOYMENT", "gpt-4")
AZURE_OPENAI_API_VERSION = os.environ.get("AZURE_OPENAI_API_VERSION", "2024-02-15-preview")

# Inicializar cliente OpenAI
client = AzureOpenAI(
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION
)


def process_paragraph_text(text: str) -> str:
    """Processa um parágrafo usando Azure OpenAI."""
    if not text or len(text.strip()) == 0:
        return text
    
    try:
        system_prompt = """Você é um corretor ortográfico profissional em português.
Sua tarefa é:
1. Corrigir todos os erros ortográficos e gramaticais
2. Eliminar redundâncias e repetições desnecessárias
3. Manter o significado e o estilo original do texto
4. Retornar APENAS o texto corrigido, sem explicações

IMPORTANTE: Retorne somente o texto corrigido."""

        response = client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Corrija este texto:\n\n{text}"}
            ],
            temperature=0.3,
            max_tokens=4000
        )
        
        return response.choices[0].message.content.strip()
        
    except Exception as e:
        logging.error(f"Erro ao processar parágrafo: {str(e)}")
        return text


def process_word_document(file_content: bytes) -> bytes:
    """Processa documento Word completo."""
    doc_stream = io.BytesIO(file_content)
    doc = Document(doc_stream)
    
    logging.info(f"Processando documento com {len(doc.paragraphs)} parágrafos")
    
    paragraphs_processed = 0
    
    # Processar parágrafos
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            try:
                original_text = paragraph.text
                corrected_text = process_paragraph_text(original_text)
                
                if corrected_text != original_text:
                    for run in paragraph.runs:
                        run.text = ""
                    if paragraph.runs:
                        paragraph.runs[0].text = corrected_text
                    else:
                        paragraph.add_run(corrected_text)
                    paragraphs_processed += 1
                    
            except Exception as e:
                logging.error(f"Erro ao processar parágrafo: {str(e)}")
                continue
    
    # Processar tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        try:
                            original_text = paragraph.text
                            corrected_text = process_paragraph_text(original_text)
                            
                            if corrected_text != original_text:
                                for run in paragraph.runs:
                                    run.text = ""
                                if paragraph.runs:
                                    paragraph.runs[0].text = corrected_text
                                else:
                                    paragraph.add_run(corrected_text)
                                    
                        except Exception as e:
                            logging.error(f"Erro ao processar célula: {str(e)}")
                            continue
    
    logging.info(f"Total de parágrafos corrigidos: {paragraphs_processed}")
    
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    
    return output_stream.getvalue()


# Blob Storage Trigger Function
# NOTA: Requer binding de Blob Storage configurado
def blob_trigger_correct_document(inputBlob: func.InputStream, 
                                   outputBlob: func.Out[bytes]) -> None:
    """
    Azure Function com Blob Trigger para processamento automático.
    
    Trigger: Blob Upload no container 'input-documents'
    Output: Blob no container 'corrected-documents'
    
    Para usar esta função, adicione ao function_app.py:
    
    @app.blob_trigger(arg_name="inputBlob", 
                      path="input-documents/{name}",
                      connection="AzureWebJobsStorage")
    @app.blob_output(arg_name="outputBlob",
                     path="corrected-documents/{name}",
                     connection="AzureWebJobsStorage")
    def blob_correct_document(inputBlob: func.InputStream, 
                              outputBlob: func.Out[bytes]) -> None:
        ...
    """
    
    logging.info(f"Processando blob: {inputBlob.name}")
    logging.info(f"Tamanho: {inputBlob.length} bytes")
    
    # Verificar extensão
    if not inputBlob.name.lower().endswith('.docx'):
        logging.warning(f"Arquivo ignorado (não é .docx): {inputBlob.name}")
        return
    
    try:
        # Ler conteúdo do blob
        file_content = inputBlob.read()
        
        # Processar documento
        corrected_content = process_word_document(file_content)
        
        # Escrever no blob de saída
        outputBlob.set(corrected_content)
        
        logging.info(f"✅ Documento processado: {inputBlob.name}")
        logging.info(f"   Tamanho original: {len(file_content)} bytes")
        logging.info(f"   Tamanho corrigido: {len(corrected_content)} bytes")
        
    except Exception as e:
        logging.error(f"❌ Erro ao processar {inputBlob.name}: {str(e)}", exc_info=True)
        raise


# Instruções de uso no README:
"""
## Configuração do Blob Storage Trigger

### 1. Criar Storage Account e Containers

```bash
# Criar containers
az storage container create --name input-documents --account-name stwordcorrection
az storage container create --name corrected-documents --account-name stwordcorrection
```

### 2. Adicionar função ao function_app.py

Adicione este código ao function_app.py:

```python
@app.blob_trigger(arg_name="inputBlob", 
                  path="input-documents/{name}",
                  connection="AzureWebJobsStorage")
@app.blob_output(arg_name="outputBlob",
                 path="corrected-documents/{name}",
                 connection="AzureWebJobsStorage")
def blob_correct_document(inputBlob: func.InputStream, 
                          outputBlob: func.Out[bytes]) -> None:
    from blob_trigger_function import process_word_document
    import logging
    
    logging.info(f"Processando: {inputBlob.name}")
    
    if not inputBlob.name.lower().endswith('.docx'):
        logging.warning(f"Ignorando: {inputBlob.name}")
        return
    
    file_content = inputBlob.read()
    corrected_content = process_word_document(file_content)
    outputBlob.set(corrected_content)
    
    logging.info(f"✅ Concluído: {inputBlob.name}")
```

### 3. Fazer Upload de Documentos

```bash
# Via Azure CLI
az storage blob upload \
  --account-name stwordcorrection \
  --container-name input-documents \
  --name documento.docx \
  --file documento.docx

# Via Azure Portal
# Navegue até o container e faça upload manualmente
```

### 4. Baixar Documentos Corrigidos

```bash
az storage blob download \
  --account-name stwordcorrection \
  --container-name corrected-documents \
  --name documento.docx \
  --file documento_corrigido.docx
```

## Vantagens do Blob Trigger

- ✅ Processamento automático quando documentos são carregados
- ✅ Suporte para processamento em lote
- ✅ Integração nativa com Azure Storage
- ✅ Escalabilidade automática
- ✅ Histórico de documentos processados
"""
