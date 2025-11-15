"""
Script de teste para a Azure Function de correção de documentos Word.

Este script cria um documento Word de exemplo com erros ortográficos,
envia para a função e valida o resultado.
"""

import requests
import io
from docx import Document


def create_test_document():
    """
    Cria um documento Word de teste com erros ortográficos propositais.
    """
    doc = Document()
    
    # Título
    doc.add_heading('Relatório de Teste - Correção Ortográfica', 0)
    
    # Parágrafo com erros
    doc.add_heading('Introdução', level=1)
    p1 = doc.add_paragraph(
        'Este é um documento de teste com varios erros ortograficos e gramaticais. '
        'O objetivo é verificar se a função de correção esta funcionando corretamente. '
        'Temos palavras escritas de forma incorreta como: occorrer, desenvolver, necessario.'
    )
    
    # Parágrafo com redundâncias
    doc.add_heading('Descrição do Problema', level=1)
    p2 = doc.add_paragraph(
        'O problema que estamos enfrentando é um problema muito sério. '
        'Este problema precisa ser resolvido urgentemente de forma urgente. '
        'A solução do problema vai resolver o problema de vez.'
    )
    
    # Tabela com erros
    doc.add_heading('Resultados', level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Cabeçalho
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Criterio'
    hdr_cells[1].text = 'Descrição'
    
    # Dados com erros
    row1 = table.rows[1].cells
    row1[0].text = 'Precisão'
    row1[1].text = 'A precisão do sistema esta muito boa e precisa'
    
    row2 = table.rows[2].cells
    row2[0].text = 'Velocidade'
    row2[1].text = 'O sistema é rapido e veloz na execução das tarefas'
    
    # Conclusão
    doc.add_heading('Conclusão', level=1)
    doc.add_paragraph(
        'Em conclusão, podemos concluir que a conclusão deste teste mostra que '
        'o sistema funciona funcionalmente de forma funcional.'
    )
    
    # Salvar em memória
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output.getvalue()


def test_local_function(endpoint="http://localhost:7071"):
    """
    Testa a função localmente.
    
    Args:
        endpoint: URL base da função (default: localhost)
    """
    print("=" * 80)
    print("TESTE DA AZURE FUNCTION - CORREÇÃO DE DOCUMENTOS WORD")
    print("=" * 80)
    
    # 1. Testar health check
    print("\n1. Testando Health Check...")
    try:
        health_response = requests.get(f"{endpoint}/api/health")
        print(f"   Status: {health_response.status_code}")
        print(f"   Resposta: {health_response.json()}")
        
        if health_response.status_code != 200:
            print("   ❌ ERRO: Health check falhou!")
            return
        print("   ✅ Health check OK!")
        
    except Exception as e:
        print(f"   ❌ ERRO: Não foi possível conectar à função: {e}")
        print("   💡 Certifique-se de que a função está rodando: func start")
        return
    
    # 2. Criar documento de teste
    print("\n2. Criando documento de teste com erros...")
    try:
        test_doc_content = create_test_document()
        print(f"   ✅ Documento criado ({len(test_doc_content)} bytes)")
    except Exception as e:
        print(f"   ❌ ERRO ao criar documento: {e}")
        return
    
    # 3. Enviar para correção
    print("\n3. Enviando documento para correção...")
    try:
        files = {"file": ("test_document.docx", test_doc_content, 
                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        
        response = requests.post(f"{endpoint}/api/correct-document", files=files)
        
        print(f"   Status: {response.status_code}")
        
        if response.status_code != 200:
            print(f"   ❌ ERRO: {response.text}")
            return
        
        corrected_content = response.content
        print(f"   ✅ Documento corrigido recebido ({len(corrected_content)} bytes)")
        
    except Exception as e:
        print(f"   ❌ ERRO ao enviar documento: {e}")
        return
    
    # 4. Salvar resultado
    print("\n4. Salvando documentos...")
    try:
        # Salvar original
        with open("test_original.docx", "wb") as f:
            f.write(test_doc_content)
        print("   ✅ Original salvo: test_original.docx")
        
        # Salvar corrigido
        with open("test_corrigido.docx", "wb") as f:
            f.write(corrected_content)
        print("   ✅ Corrigido salvo: test_corrigido.docx")
        
    except Exception as e:
        print(f"   ❌ ERRO ao salvar arquivos: {e}")
        return
    
    # 5. Comparar resultados
    print("\n5. Comparando resultados...")
    try:
        original_doc = Document(io.BytesIO(test_doc_content))
        corrected_doc = Document(io.BytesIO(corrected_content))
        
        original_paragraphs = [p.text for p in original_doc.paragraphs if p.text.strip()]
        corrected_paragraphs = [p.text for p in corrected_doc.paragraphs if p.text.strip()]
        
        print(f"   Parágrafos no original: {len(original_paragraphs)}")
        print(f"   Parágrafos no corrigido: {len(corrected_paragraphs)}")
        
        # Mostrar algumas diferenças
        print("\n   📝 Exemplos de correções:")
        for i, (orig, corr) in enumerate(zip(original_paragraphs[:3], corrected_paragraphs[:3])):
            if orig != corr:
                print(f"\n   Parágrafo {i+1}:")
                print(f"   ANTES: {orig[:100]}...")
                print(f"   DEPOIS: {corr[:100]}...")
        
    except Exception as e:
        print(f"   ⚠️ Aviso: Não foi possível comparar documentos: {e}")
    
    print("\n" + "=" * 80)
    print("✅ TESTE CONCLUÍDO COM SUCESSO!")
    print("=" * 80)
    print("\n📂 Arquivos gerados:")
    print("   - test_original.docx  (documento com erros)")
    print("   - test_corrigido.docx (documento corrigido)")
    print("\n💡 Abra os arquivos no Word para verificar as correções!")


if __name__ == "__main__":
    import sys
    
    # Permitir especificar endpoint customizado
    endpoint = sys.argv[1] if len(sys.argv) > 1 else "http://localhost:7071"
    
    print(f"\n🎯 Testando endpoint: {endpoint}")
    test_local_function(endpoint)
